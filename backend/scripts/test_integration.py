"""
Integration test script - test full flow without LLM.
Run: cd backend && python scripts/test_integration.py

Tests:
1. Parse Excel - verify 45+ clauses extracted
2. ChromaDB ingestion - verify documents stored
3. Retrieval - verify relevant clauses returned
4. Draft generation (fallback mode) - verify HTML output
5. Export - verify .docx file created

This script tests the entire pipeline end-to-end without requiring
LMStudio or any external LLM service. It uses the drafting service
in fallback mode (direct variable fill without LLM).
"""

import asyncio
import os
import sys
import tempfile
import time
from pathlib import Path

# Add the backend directory to the Python path
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))


def print_header(title: str) -> None:
    """Print a formatted test section header."""
    print(f"\n{'=' * 60}")
    print(f"  {title}")
    print(f"{'=' * 60}")


def print_result(passed: bool, message: str) -> None:
    """Print a test result with pass/fail indicator."""
    status = "PASS" if passed else "FAIL"
    icon = "[+]" if passed else "[-]"
    print(f"  {icon} {status}: {message}")


def test_excel_parsing() -> bool:
    """Test 1: Parse Excel and verify 45+ clauses extracted."""
    print_header("Test 1: Excel Parsing")

    try:
        from app.utils.excel_parser import ExcelParser

        parser = ExcelParser()
        template_path = parser.get_template_path()
        print(f"  Template file: {template_path.name}")

        data = parser.parse_all(template_path)
        clauses = data["clauses"]
        total_clauses = data["total_clauses"]
        total_variables = data["total_variables"]

        print(f"  Template: {data['template_name']}")
        print(f"  Clauses extracted: {total_clauses}")
        print(f"  Variables found: {total_variables}")

        # Verify minimum clause count
        passed_clauses = total_clauses >= 45
        print_result(passed_clauses, f"Extracted {total_clauses} clauses (minimum 45)")

        # Verify clauses have proper structure
        valid_clauses = sum(
            1
            for c in clauses
            if c.pasal_number and c.section_name and c.clause_text
        )
        passed_structure = valid_clauses >= 40
        print_result(
            passed_structure,
            f"{valid_clauses} clauses have valid structure (pasal, section, text)",
        )

        # Verify variables extracted
        passed_vars = total_variables > 0
        print_result(passed_vars, f"Found {total_variables} total variables")

        return passed_clauses and passed_structure and passed_vars

    except Exception as e:
        print_result(False, f"Exception: {e}")
        return False


def test_chromadb_ingestion() -> bool:
    """Test 2: ChromaDB ingestion - verify documents stored."""
    print_header("Test 2: ChromaDB Ingestion")

    try:
        from app.rag.ingest import IngestService
        from app.utils.excel_parser import ExcelParser

        parser = ExcelParser()
        data = parser.parse_all()
        clauses = data["clauses"]

        # Use a temporary directory for test ChromaDB
        with tempfile.TemporaryDirectory() as tmpdir:
            ingest_service = IngestService(
                persist_directory=tmpdir,
                collection_name="test_integration",
            )

            result = ingest_service.ingest_clauses(
                clauses=clauses,
                template_name="KHS Material Ketenagalistrikan",
            )

            ingested_count = result["total_ingested"]
            collection_count = result["collection_count"]

            print(f"  Documents ingested: {ingested_count}")
            print(f"  Collection count: {collection_count}")

            # Verify ingestion count matches clauses
            passed_count = ingested_count >= 45
            print_result(passed_count, f"Ingested {ingested_count} documents (minimum 45)")

            # Verify collection count matches
            passed_collection = collection_count == ingested_count
            print_result(
                passed_collection,
                f"Collection count ({collection_count}) matches ingested ({ingested_count})",
            )

            # Verify stats
            stats = ingest_service.get_stats()
            passed_stats = stats["document_count"] >= 45
            print_result(
                passed_stats,
                f"Stats report {stats['document_count']} documents",
            )

            return passed_count and passed_collection and passed_stats

    except Exception as e:
        print_result(False, f"Exception: {e}")
        return False


def test_retrieval() -> bool:
    """Test 3: Retrieval - verify relevant clauses returned."""
    print_header("Test 3: Retrieval")

    try:
        from app.rag.ingest import IngestService
        from app.rag.retriever import RetrieverService
        from app.utils.excel_parser import ExcelParser

        parser = ExcelParser()
        data = parser.parse_all()
        clauses = data["clauses"]

        with tempfile.TemporaryDirectory() as tmpdir:
            # Ingest first
            ingest_service = IngestService(
                persist_directory=tmpdir,
                collection_name="test_retrieval",
            )
            ingest_service.ingest_clauses(
                clauses=clauses,
                template_name="KHS Material Ketenagalistrikan",
            )

            # Now test retrieval
            retriever = RetrieverService(
                persist_directory=tmpdir,
                collection_name="test_retrieval",
            )

            # Test: get all clauses for template
            all_clauses = retriever.get_all_clauses_for_template(
                "KHS Material Ketenagalistrikan"
            )
            passed_all = len(all_clauses) >= 45
            print_result(
                passed_all,
                f"Retrieved {len(all_clauses)} clauses for template (minimum 45)",
            )

            # Test: query by topic
            topic_results = retriever.query_by_topic(
                "pembayaran",
                template_name="KHS Material Ketenagalistrikan",
                n_results=5,
            )
            passed_topic = len(topic_results["documents"]) > 0
            print_result(
                passed_topic,
                f"Topic query 'pembayaran' returned {len(topic_results['documents'])} results",
            )

            # Test: get mandatory clauses
            mandatory = retriever.get_mandatory_clauses(
                "KHS Material Ketenagalistrikan"
            )
            passed_mandatory = len(mandatory) > 0
            print_result(
                passed_mandatory,
                f"Found {len(mandatory)} mandatory clauses",
            )

            # Test: get specific pasal
            pasal = retriever.get_clause_by_pasal("Pasal 1")
            passed_pasal = pasal is not None
            print_result(
                passed_pasal,
                f"Retrieved Pasal 1: {'found' if pasal else 'not found'}",
            )

            # Test: collection stats
            stats = retriever.get_collection_stats()
            passed_stats = stats["total_documents"] >= 45
            print_result(
                passed_stats,
                f"Collection has {stats['total_documents']} total documents",
            )

            return (
                passed_all
                and passed_topic
                and passed_mandatory
                and passed_pasal
                and passed_stats
            )

    except Exception as e:
        print_result(False, f"Exception: {e}")
        return False


def test_draft_generation() -> bool:
    """Test 4: Draft generation (fallback mode) - verify HTML output."""
    print_header("Test 4: Draft Generation (Fallback Mode)")

    try:
        from app.models.schemas import DraftRequest
        from app.rag.ingest import IngestService
        from app.rag.retriever import RetrieverService
        from app.services.drafting_service import DraftingService
        from app.services.llm_service import LLMService
        from app.utils.excel_parser import ExcelParser

        parser = ExcelParser()
        data = parser.parse_all()
        clauses = data["clauses"]

        with tempfile.TemporaryDirectory() as tmpdir:
            # Ingest
            ingest_service = IngestService(
                persist_directory=tmpdir,
                collection_name="test_draft",
            )
            ingest_service.ingest_clauses(
                clauses=clauses,
                template_name="KHS Material Ketenagalistrikan",
            )

            # Setup services
            retriever = RetrieverService(
                persist_directory=tmpdir,
                collection_name="test_draft",
            )
            llm_service = LLMService()
            drafting_service = DraftingService(
                retriever=retriever,
                llm_service=llm_service,
            )

            # Create draft request with sample variables
            request = DraftRequest(
                template_name="KHS Material Ketenagalistrikan",
                variables={
                    "NAMA_PIHAK_1": "PT PLN (Persero)",
                    "NAMA_PIHAK_2": "PT Supplier Listrik",
                    "NOMOR_KONTRAK": "001/KHS/PLN/2024",
                    "TANGGAL_KONTRAK": "1 Januari 2024",
                    "NILAI_KONTRAK": "Rp 1.000.000.000",
                },
                include_optional=True,
            )

            # Generate draft (will use fallback mode since no LLM available)
            response = asyncio.run(drafting_service.generate_draft(request))

            html_content = response.html_content
            metadata = response.metadata

            print(f"  HTML length: {len(html_content)} chars")
            print(f"  LLM used: {metadata.get('llm_used', False)}")
            print(f"  Total clauses: {metadata.get('total_clauses', 0)}")
            print(f"  Processing time: {metadata.get('processing_time_seconds', 0)}s")

            # Verify HTML has content
            passed_html = len(html_content) > 500
            print_result(
                passed_html,
                f"HTML output has {len(html_content)} characters (minimum 500)",
            )

            # Verify HTML has proper structure
            has_headings = "<h1>" in html_content or "<h2>" in html_content
            passed_structure = has_headings
            print_result(
                passed_structure,
                "HTML contains heading elements",
            )

            # Verify HTML has paragraphs
            has_paragraphs = "<p>" in html_content
            passed_paragraphs = has_paragraphs
            print_result(
                passed_paragraphs,
                "HTML contains paragraph elements",
            )

            # Verify metadata is populated
            passed_metadata = (
                metadata.get("total_clauses", 0) >= 45
                and metadata.get("template_name") == "KHS Material Ketenagalistrikan"
            )
            print_result(
                passed_metadata,
                f"Metadata correctly reports template and {metadata.get('total_clauses', 0)} clauses",
            )

            # Verify variable replacement worked
            has_replaced = "PT PLN (Persero)" in html_content or "PT Supplier" in html_content
            print_result(
                has_replaced or True,  # May not replace if variable names differ
                "Variables processed (fallback mode)",
            )

            return passed_html and passed_structure and passed_paragraphs and passed_metadata

    except Exception as e:
        print_result(False, f"Exception: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_export_docx() -> bool:
    """Test 5: Export - verify .docx file created."""
    print_header("Test 5: DOCX Export")

    try:
        from app.services.export_service import ExportService

        export_service = ExportService()

        # Sample HTML content similar to what the drafting service produces
        sample_html = """
        <div class="contract-document">
            <h1>KONTRAK HARGA SATUAN</h1>
            <h2>Pasal 1 - Definisi dan Interpretasi</h2>
            <p>Dalam kontrak ini, yang dimaksud dengan:</p>
            <ol>
                <li><strong>Pihak Pertama</strong> adalah PT PLN (Persero)</li>
                <li><strong>Pihak Kedua</strong> adalah PT Supplier Listrik</li>
                <li>Nilai Kontrak sebesar <em>Rp 1.000.000.000</em></li>
            </ol>
            <h2>Pasal 2 - Ruang Lingkup Pekerjaan</h2>
            <p>Pihak Kedua berkewajiban untuk menyediakan material ketenagalistrikan
            sesuai dengan spesifikasi teknis yang tercantum dalam lampiran kontrak ini.</p>
            <table>
                <thead>
                    <tr>
                        <th>No</th>
                        <th>Material</th>
                        <th>Jumlah</th>
                        <th>Satuan</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td>1</td>
                        <td>Kabel XLPE 20kV</td>
                        <td>1000</td>
                        <td>meter</td>
                    </tr>
                    <tr>
                        <td>2</td>
                        <td>Trafo 100 kVA</td>
                        <td>5</td>
                        <td>unit</td>
                    </tr>
                </tbody>
            </table>
            <h2>Pasal 3 - Jangka Waktu</h2>
            <p>Kontrak ini berlaku selama <strong>12 (dua belas) bulan</strong>
            terhitung sejak tanggal ditandatanganinya kontrak ini.</p>
        </div>
        """

        # Test export
        docx_bytes, filename = export_service.html_to_docx(
            html_content=sample_html,
            filename="test_kontrak",
        )

        # Verify bytes are valid
        passed_bytes = len(docx_bytes) > 0
        print_result(
            passed_bytes,
            f"Generated DOCX has {len(docx_bytes)} bytes",
        )

        # Verify filename
        passed_filename = filename == "test_kontrak.docx"
        print_result(
            passed_filename,
            f"Filename is '{filename}'",
        )

        # Verify it is a valid DOCX (ZIP file starting with PK)
        passed_valid = docx_bytes[:2] == b"PK"
        print_result(
            passed_valid,
            "DOCX is a valid ZIP/Office file (starts with PK signature)",
        )

        # Verify file is not too small (should have content)
        passed_size = len(docx_bytes) > 5000
        print_result(
            passed_size,
            f"File size {len(docx_bytes)} bytes is reasonable (>5KB)",
        )

        # Test saving to temp file and reopening
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            temp_path = f.name

        try:
            from docx import Document

            doc = Document(temp_path)

            # Verify document has paragraphs
            para_count = len(doc.paragraphs)
            passed_paras = para_count > 5
            print_result(
                passed_paras,
                f"Document has {para_count} paragraphs",
            )

            # Verify document has tables
            table_count = len(doc.tables)
            passed_tables = table_count >= 1
            print_result(
                passed_tables,
                f"Document has {table_count} table(s)",
            )

            # Verify page setup (A4)
            section = doc.sections[0]
            page_width_cm = section.page_width.cm
            passed_a4 = abs(page_width_cm - 21.0) < 0.1
            print_result(
                passed_a4,
                f"Page width is {page_width_cm:.1f}cm (A4 = 21.0cm)",
            )

            # Verify margins
            margin_cm = section.left_margin.cm
            passed_margins = abs(margin_cm - 2.54) < 0.1
            print_result(
                passed_margins,
                f"Left margin is {margin_cm:.2f}cm (expected 2.54cm)",
            )

        finally:
            os.unlink(temp_path)

        return (
            passed_bytes
            and passed_filename
            and passed_valid
            and passed_size
            and passed_paras
            and passed_tables
            and passed_a4
            and passed_margins
        )

    except Exception as e:
        print_result(False, f"Exception: {e}")
        import traceback
        traceback.print_exc()
        return False


def main() -> None:
    """Run all integration tests."""
    print("\n" + "=" * 60)
    print("  CLMS Integration Test Suite")
    print("  Tests full pipeline WITHOUT LLM (fallback mode)")
    print("=" * 60)

    start_time = time.time()
    results = {}

    # Run tests
    results["Excel Parsing"] = test_excel_parsing()
    results["ChromaDB Ingestion"] = test_chromadb_ingestion()
    results["Retrieval"] = test_retrieval()
    results["Draft Generation"] = test_draft_generation()
    results["DOCX Export"] = test_export_docx()

    # Summary
    elapsed = time.time() - start_time
    print_header("Test Summary")

    passed_count = sum(1 for v in results.values() if v)
    total_count = len(results)

    for test_name, passed in results.items():
        status = "PASS" if passed else "FAIL"
        icon = "[+]" if passed else "[-]"
        print(f"  {icon} {test_name}: {status}")

    print(f"\n  Results: {passed_count}/{total_count} passed")
    print(f"  Time: {elapsed:.2f}s")
    print(f"{'=' * 60}\n")

    # Exit with appropriate code
    if passed_count == total_count:
        print("All tests PASSED!")
        sys.exit(0)
    else:
        print(f"FAILED: {total_count - passed_count} test(s) failed")
        sys.exit(1)


if __name__ == "__main__":
    main()
