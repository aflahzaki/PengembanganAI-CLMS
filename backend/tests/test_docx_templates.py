"""Tests for DOCX Template Library service and API endpoints."""

import io
import os
import shutil
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.docx_template_service import DocxTemplateService


# --- Fixtures ---


@pytest.fixture
def temp_templates_dir():
    """Create a temporary directory for test templates."""
    tmpdir = tempfile.mkdtemp()
    yield Path(tmpdir)
    shutil.rmtree(tmpdir)


@pytest.fixture
def service(temp_templates_dir):
    """Create a DocxTemplateService with a temporary directory."""
    return DocxTemplateService(templates_dir=temp_templates_dir)


@pytest.fixture
def sample_docx(temp_templates_dir) -> Path:
    """Create a sample DOCX file for testing."""
    filepath = temp_templates_dir / "20241007 PLN - BPR - Template Kontrak Lumsum.docx"
    doc = Document()

    # Add heading
    doc.add_heading("SURAT PERJANJIAN", level=1)

    # Add paragraph with dynamic variable (ellipsis)
    doc.add_paragraph("Nomor: [...]")

    # Add paragraph with editable variable (short descriptive)
    doc.add_paragraph("Yang bertanda tangan di bawah ini [diisi dengan nama pekerjaan]")

    # Add paragraph with instruction variable (>60 chars)
    long_text = (
        "[Klausul dibawah dapat dipilih sesuai dengan kebutuhan pelaksanaan "
        "pekerjaan dalam Perjanjian ini]"
    )
    doc.add_paragraph(f"Catatan: {long_text}")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "No"
    table.rows[0].cells[1].text = "Deskripsi"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "[Masukan Deskripsi]"

    doc.save(str(filepath))
    return filepath


@pytest.fixture
def client():
    """Create a test client for the FastAPI app."""
    return TestClient(app)


# --- Unit Tests for DocxTemplateService ---


class TestSlugify:
    """Tests for the _slugify method."""

    def test_standard_filename(self, service):
        result = service._slugify(
            "20241007 PLN - BPR - Template Kontrak Lumsum.docx"
        )
        assert result == "template-kontrak-lumsum"

    def test_filename_with_multiple_words(self, service):
        result = service._slugify(
            "20241007 PLN - BPR - Template Kontrak Biaya Plus Imbalan.docx"
        )
        assert result == "template-kontrak-biaya-plus-imbalan"

    def test_simple_filename(self, service):
        result = service._slugify("My Custom Template.docx")
        assert result == "my-custom-template"


class TestClassifyVariable:
    """Tests for the _classify_variable method."""

    def test_ellipsis_dots(self, service):
        assert service._classify_variable("...") == "dynamic"

    def test_ellipsis_unicode(self, service):
        assert service._classify_variable("\u2026") == "dynamic"

    def test_ellipsis_in_parens(self, service):
        assert service._classify_variable("(\u2026)") == "dynamic"

    def test_ellipsis_dots_in_parens(self, service):
        assert service._classify_variable("(...)") == "dynamic"

    def test_short_descriptive(self, service):
        assert service._classify_variable("diisi dengan nama pekerjaan") == "editable"

    def test_short_descriptive_with_symbols(self, service):
        assert service._classify_variable("Barang/Jasa") == "editable"

    def test_long_instruction(self, service):
        long_text = (
            "Klausul dibawah dapat dipilih sesuai dengan kebutuhan pelaksanaan "
            "pekerjaan dalam Perjanjian"
        )
        assert service._classify_variable(long_text) == "instruction"

    def test_exactly_60_chars_is_editable(self, service):
        text = "a" * 60
        assert service._classify_variable(text) == "editable"

    def test_61_chars_is_instruction(self, service):
        text = "a" * 61
        assert service._classify_variable(text) == "instruction"


class TestListTemplates:
    """Tests for listing templates."""

    def test_empty_directory(self, service):
        result = service.list_templates()
        assert result == []

    def test_list_with_sample(self, service, sample_docx):
        result = service.list_templates()
        assert len(result) == 1
        template = result[0]
        assert template["id"] == "template-kontrak-lumsum"
        assert template["name"] == "Template Kontrak Lumsum"
        assert template["variable_count"] > 0
        assert template["file_size_bytes"] > 0

    def test_list_returns_variables(self, service, sample_docx):
        result = service.list_templates()
        template = result[0]
        var_types = [v["type"] for v in template["variables"]]
        assert "dynamic" in var_types
        assert "editable" in var_types


class TestParseToHtml:
    """Tests for HTML parsing with highlighting."""

    def test_not_found(self, service):
        result = service.parse_to_html("nonexistent-template")
        assert result is None

    def test_parse_sample(self, service, sample_docx):
        result = service.parse_to_html("template-kontrak-lumsum")
        assert result is not None
        assert result["id"] == "template-kontrak-lumsum"
        assert result["name"] == "Template Kontrak Lumsum"
        assert "<h1>" in result["html_content"]

    def test_dynamic_highlight(self, service, sample_docx):
        result = service.parse_to_html("template-kontrak-lumsum")
        html = result["html_content"]
        # Dynamic variable should have bright yellow highlight
        assert 'background-color: #FFEB3B' in html

    def test_editable_highlight(self, service, sample_docx):
        result = service.parse_to_html("template-kontrak-lumsum")
        html = result["html_content"]
        # Editable variable should have light yellow highlight
        assert 'background-color: #FFF9C4' in html

    def test_instruction_styling(self, service, sample_docx):
        result = service.parse_to_html("template-kontrak-lumsum")
        html = result["html_content"]
        # Instruction variable should have gray italic
        assert 'color: gray; font-style: italic' in html

    def test_contains_table(self, service, sample_docx):
        result = service.parse_to_html("template-kontrak-lumsum")
        html = result["html_content"]
        assert "<table>" in html
        assert "<td>" in html

    def test_split_run_brackets_highlighted(self, service, temp_templates_dir):
        """Test that brackets split across runs are still highlighted."""
        filepath = temp_templates_dir / "Split Run Test.docx"
        doc = Document()
        para = doc.add_paragraph()
        # Simulate Word splitting a bracket across multiple runs
        run1 = para.add_run("[diisi")
        run1.bold = True
        run2 = para.add_run(" dengan nama]")
        run2.bold = False
        doc.save(str(filepath))

        result = service.parse_to_html("split-run-test")
        assert result is not None
        html = result["html_content"]
        # The bracket content should be matched and highlighted
        assert "background-color: #FFF9C4" in html


class TestUploadAndDelete:
    """Tests for upload and delete operations."""

    def test_upload_valid_docx(self, service, temp_templates_dir):
        # Create a valid docx in memory
        doc = Document()
        doc.add_paragraph("Test content with [Variable]")
        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        result = service.save_uploaded_template("Test Template.docx", content)
        assert result["id"] == "test-template"
        assert result["filename"] == "Test Template.docx"
        assert result["variable_count"] == 1

        # Verify file exists
        assert (temp_templates_dir / "Test Template.docx").exists()

    def test_upload_non_docx_raises(self, service):
        with pytest.raises(ValueError, match="Only .docx files"):
            service.save_uploaded_template("test.txt", b"content")

    def test_upload_path_traversal_blocked(self, service, temp_templates_dir):
        """Test that path traversal in filename is blocked."""
        # Create a valid docx in memory
        doc = Document()
        doc.add_paragraph("Traversal test")
        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        result = service.save_uploaded_template("../../evil.docx", content)
        # Should strip directory components and save only the basename
        assert result["filename"] == "evil.docx"
        assert (temp_templates_dir / "evil.docx").exists()
        # Ensure it did NOT write outside the templates dir
        assert not (temp_templates_dir.parent.parent / "evil.docx").exists()

    def test_upload_invalid_docx_content_rejected(self, service):
        """Test that invalid DOCX content is rejected."""
        with pytest.raises(ValueError, match="not a valid DOCX"):
            service.save_uploaded_template("fake.docx", b"not a real docx file")

    def test_delete_existing(self, service, sample_docx):
        result = service.delete_template("template-kontrak-lumsum")
        assert result is True
        assert not sample_docx.exists()

    def test_delete_nonexistent(self, service):
        result = service.delete_template("nonexistent")
        assert result is False


# --- API Integration Tests ---


class TestDocxTemplateAPI:
    """Tests for the DOCX template API endpoints."""

    def test_list_templates_endpoint(self, client):
        """Test that the list endpoint returns a valid response."""
        response = client.get("/api/templates/docx")
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)
        # Should have the 8 templates we copied
        assert len(data) == 8

    def test_get_template_html_endpoint(self, client):
        """Test getting HTML for a known template."""
        response = client.get("/api/templates/docx/template-kontrak-lumsum/html")
        assert response.status_code == 200
        data = response.json()
        assert data["id"] == "template-kontrak-lumsum"
        assert "html_content" in data
        assert len(data["html_content"]) > 0

    def test_get_template_html_not_found(self, client):
        """Test 404 for nonexistent template."""
        response = client.get("/api/templates/docx/nonexistent-template/html")
        assert response.status_code == 404

    def test_upload_template_endpoint(self, client):
        """Test uploading a new template."""
        # Create a minimal docx file
        doc = Document()
        doc.add_paragraph("Upload test [Variable]")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = client.post(
            "/api/templates/docx/upload",
            files={"file": ("Test Upload.docx", buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["id"] == "test-upload"
        assert data["filename"] == "Test Upload.docx"

        # Clean up: delete the uploaded file
        client.delete(f"/api/templates/docx/{data['id']}")

    def test_upload_non_docx_rejected(self, client):
        """Test that non-docx files are rejected."""
        buffer = io.BytesIO(b"not a docx file")
        response = client.post(
            "/api/templates/docx/upload",
            files={"file": ("test.txt", buffer, "text/plain")},
        )
        assert response.status_code == 400

    def test_upload_invalid_docx_content_rejected(self, client):
        """Test that files with .docx extension but invalid content are rejected."""
        buffer = io.BytesIO(b"this is not a real docx file")
        response = client.post(
            "/api/templates/docx/upload",
            files={"file": ("fake.docx", buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 400
        assert "not a valid DOCX" in response.json()["detail"]

    def test_delete_template_endpoint(self, client):
        """Test deleting a template (upload then delete)."""
        # Upload a template first
        doc = Document()
        doc.add_paragraph("Delete test")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        upload_response = client.post(
            "/api/templates/docx/upload",
            files={"file": ("Delete Me.docx", buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        template_id = upload_response.json()["id"]

        # Delete it
        response = client.delete(f"/api/templates/docx/{template_id}")
        assert response.status_code == 204

        # Verify it's gone
        response = client.get(f"/api/templates/docx/{template_id}/html")
        assert response.status_code == 404

    def test_delete_nonexistent_template(self, client):
        """Test 404 when deleting nonexistent template."""
        response = client.delete("/api/templates/docx/nonexistent-id")
        assert response.status_code == 404

    def test_template_variables_classification(self, client):
        """Test that variables are properly classified in list response."""
        response = client.get("/api/templates/docx")
        assert response.status_code == 200
        data = response.json()

        # At least one template should have variables
        templates_with_vars = [t for t in data if t["variable_count"] > 0]
        assert len(templates_with_vars) > 0

        # Check variable types are valid
        for template in templates_with_vars:
            for var in template["variables"]:
                assert var["type"] in ("dynamic", "editable", "instruction")
                assert "name" in var
                assert "full_text" in var
