"""Tests for image handling in the export service (HTML to DOCX conversion)."""

import base64
import io
import os
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from app.services.export_service import ExportService


def _create_minimal_png() -> bytes:
    """Create a minimal valid 1x1 pixel PNG image for testing.

    Returns:
        bytes: Valid PNG file content.
    """
    # Minimal 1x1 white pixel PNG
    import struct
    import zlib

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        raw = chunk_type + data
        return struct.pack(">I", len(data)) + raw + struct.pack(">I", zlib.crc32(raw) & 0xFFFFFFFF)

    # PNG signature
    signature = b"\x89PNG\r\n\x1a\n"
    # IHDR: width=1, height=1, bit_depth=8, color_type=2 (RGB)
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = _chunk(b"IHDR", ihdr_data)
    # IDAT: raw pixel data (filter byte 0 + RGB white pixel)
    raw_data = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw_data)
    idat = _chunk(b"IDAT", compressed)
    # IEND
    iend = _chunk(b"IEND", b"")

    return signature + ihdr + idat + iend


class TestExportImageBase64:
    """Tests for base64 data URI image handling in DOCX export."""

    def test_base64_image_embedded_in_docx(self):
        """Test that a base64 image in HTML is embedded in the output DOCX."""
        png_bytes = _create_minimal_png()
        b64_data = base64.b64encode(png_bytes).decode("utf-8")
        data_uri = f"data:image/png;base64,{b64_data}"

        html = f'<div><p>Before image</p><img src="{data_uri}" alt="Signature" /><p>After image</p></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html, "test_image")

        # Verify valid DOCX is produced
        assert docx_bytes is not None
        assert len(docx_bytes) > 0
        assert filename == "test_image.docx"

        # Verify can be opened as a Document
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Verify the document contains inline shapes (the image)
        # python-docx stores images as inline shapes
        has_image = False
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_image = True
                break
        assert has_image, "DOCX should contain an embedded image"

    def test_base64_image_inline_in_paragraph(self):
        """Test that an inline base64 image within a <p> tag is handled."""
        png_bytes = _create_minimal_png()
        b64_data = base64.b64encode(png_bytes).decode("utf-8")
        data_uri = f"data:image/png;base64,{b64_data}"

        html = f'<div><p>Text before <img src="{data_uri}" alt="sig" /> text after</p></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        # Verify valid DOCX
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Check that it has an image relationship
        has_image = False
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_image = True
                break
        assert has_image, "DOCX should contain an inline image from paragraph"

    def test_invalid_base64_shows_placeholder(self):
        """Test that invalid base64 data gracefully shows a placeholder."""
        html = '<div><img src="data:image/png;base64,NOT_VALID_BASE64!!!" alt="BadImage" /></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        # Should still produce a valid DOCX
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Check for placeholder text
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[BadImage]" in full_text


class TestExportImageLocalFile:
    """Tests for local file path image handling in DOCX export."""

    def test_local_uploads_image_embedded(self, tmp_path):
        """Test that a local /uploads/ image path resolves and embeds."""
        png_bytes = _create_minimal_png()

        # Create a fake uploads directory structure
        uploads_dir = tmp_path / "data" / "uploads"
        uploads_dir.mkdir(parents=True)
        image_file = uploads_dir / "signature.png"
        image_file.write_bytes(png_bytes)

        html = '<div><img src="/uploads/signature.png" alt="Tanda Tangan" /></div>'

        service = ExportService()

        # Patch the path resolution to use our tmp directory
        with patch.object(Path, "is_file", return_value=True):
            with patch.object(Path, "read_bytes", return_value=png_bytes):
                docx_bytes, filename = service.html_to_docx(html)

        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Verify image was embedded
        has_image = False
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_image = True
                break
        assert has_image, "DOCX should contain the local file image"

    def test_missing_local_file_shows_placeholder(self):
        """Test that a missing local file gracefully shows placeholder text."""
        html = '<div><img src="/uploads/nonexistent.png" alt="Missing" /></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # The placeholder text should appear
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[Missing]" in full_text

    def test_path_traversal_blocked(self):
        """Test that path traversal in image src is blocked."""
        html = '<div><img src="/uploads/../../etc/passwd" alt="Traversal" /></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

        # Should show placeholder, not the file content
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[Traversal]" in full_text


class TestExportImageNoSrc:
    """Tests for edge cases with <img> tags."""

    def test_img_without_src_shows_placeholder(self):
        """Test that an <img> tag without src shows placeholder."""
        html = '<div><img alt="NoSrc" /></div>'

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[NoSrc]" in full_text

    def test_img_without_src_or_alt_shows_default(self):
        """Test that an <img> tag with no src and no alt shows [Image]."""
        html = "<div><img /></div>"

        service = ExportService()
        docx_bytes, filename = service.html_to_docx(html)

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[Image]" in full_text
