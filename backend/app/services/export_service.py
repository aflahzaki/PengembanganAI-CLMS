"""Export service for converting HTML contracts to document formats.

Converts HTML contract content to .docx format using python-docx and
BeautifulSoup for robust HTML parsing. Produces professional legal
documents with proper formatting, headers, footers, and page setup.
"""

import base64
import io
import re
from pathlib import Path
from typing import Optional, Tuple

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement


class ExportService:
    """Service for exporting contracts to various document formats.

    Currently supports:
    - DOCX (Microsoft Word) via python-docx

    Features:
    - A4 page setup with 2.54cm margins
    - Header: "KONTRAK HARGA SATUAN - RAHASIA"
    - Footer: page numbers
    - Times New Roman 12pt body, headings bold
    - 1.5 line spacing, 6pt after paragraph
    - Table support from HTML <table> elements
    - Page break before Pasal 1 (after preamble)
    """

    def __init__(self, output_dir: Optional[str] = None):
        """Initialize export service.

        Args:
            output_dir: Optional directory for saving exported files.
        """
        self.output_dir = Path(output_dir) if output_dir else None

    def html_to_docx(
        self,
        html_content: str,
        filename: Optional[str] = None,
    ) -> Tuple[bytes, str]:
        """Convert HTML content to a DOCX document.

        Uses BeautifulSoup for robust HTML parsing and maps elements
        to Word document structures:
        - <h1> -> Heading 1 (centered)
        - <h2> -> Heading 2
        - <h3> -> Heading 3
        - <p> -> Normal paragraph with 1.5 line spacing
        - <ol>/<li> -> Numbered list items
        - <ul>/<li> -> Bullet list items
        - <strong>/<b> -> Bold text
        - <em>/<i> -> Italic text
        - <table> -> Word table

        Args:
            html_content: HTML string to convert.
            filename: Optional output filename (without extension).

        Returns:
            Tuple of (docx_bytes, filename).
        """
        document = Document()

        # Setup page (A4, margins 2.54cm)
        self._setup_page(document)

        # Setup default styles
        self._setup_styles(document)

        # Add header and footer
        self._add_header_footer(document)

        # Parse and convert HTML elements using BeautifulSoup
        self._parse_html_to_docx(html_content, document)

        # Save to bytes
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        output_filename = filename or "kontrak_draft"
        if not output_filename.endswith(".docx"):
            output_filename += ".docx"

        return buffer.getvalue(), output_filename

    def _setup_page(self, document: Document) -> None:
        """Configure A4 page setup with 2.54cm (1 inch) margins.

        Args:
            document: python-docx Document to configure.
        """
        section = document.sections[0]

        # A4 paper size
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

        # Margins: 2.54cm (1 inch) on all sides
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def _setup_styles(self, document: Document) -> None:
        """Configure default document styles.

        Sets Times New Roman 12pt for Normal style with 1.5 line spacing.

        Args:
            document: python-docx Document to configure.
        """
        # Normal style
        style = document.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Set paragraph format for Normal style
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.5

        # Heading styles - bold, Times New Roman
        for level in range(1, 4):
            heading_style_name = f"Heading {level}"
            if heading_style_name in document.styles:
                heading_style = document.styles[heading_style_name]
                heading_style.font.name = "Times New Roman"
                heading_style.font.bold = True
                if level == 1:
                    heading_style.font.size = Pt(14)
                elif level == 2:
                    heading_style.font.size = Pt(12)
                else:
                    heading_style.font.size = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
                heading_style.paragraph_format.line_spacing = 1.5

    def _add_header_footer(self, document: Document) -> None:
        """Add header and footer to the document.

        Header: "KONTRAK HARGA SATUAN - RAHASIA"
        Footer: Page number (centered)

        Args:
            document: python-docx Document to add header/footer to.
        """
        section = document.sections[0]

        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("KONTRAK HARGA SATUAN - RAHASIA")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.bold = True

        # Footer with page number
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        run2 = footer_para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._r.append(instrText)

        run3 = footer_para.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar2)

    def _parse_html_to_docx(self, html: str, document: Document) -> None:
        """Parse HTML using BeautifulSoup and add elements to the Word document.

        Handles nested tags, multi-line HTML, tables, and complex structures
        that regex-based parsing would miss.

        Args:
            html: HTML content string.
            document: python-docx Document to populate.
        """
        soup = BeautifulSoup(html, "html.parser")

        # Find the main content - could be wrapped in a div or be the root
        body = soup.find("div") or soup

        # Track whether we have passed the preamble (for page break before Pasal 1)
        self._page_break_added = False

        # Process all top-level elements
        for element in body.children:
            if isinstance(element, NavigableString):
                text = element.strip()
                if text:
                    para = document.add_paragraph()
                    self._apply_paragraph_format(para)
                    para.add_run(text)
                continue

            if isinstance(element, Tag):
                self._process_element(element, document)

    def _process_element(self, element: Tag, document: Document) -> None:
        """Process a single HTML element and add it to the document.

        Args:
            element: BeautifulSoup Tag element to process.
            document: python-docx Document to add content to.
        """
        tag_name = element.name.lower() if element.name else ""

        # Headings
        if tag_name == "h1":
            text = element.get_text(strip=True)
            para = document.add_heading(text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._apply_heading_font(para)
            return

        if tag_name == "h2":
            text = element.get_text(strip=True)
            # Add page break before first "Pasal 1" heading
            if not self._page_break_added and "pasal 1" in text.lower():
                self._page_break_added = True
                # Add page break before Pasal 1
                para = document.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)
            para = document.add_heading(text, level=2)
            self._apply_heading_font(para)
            return

        if tag_name == "h3":
            text = element.get_text(strip=True)
            para = document.add_heading(text, level=3)
            self._apply_heading_font(para)
            return

        # Paragraphs
        if tag_name == "p":
            para = document.add_paragraph()
            self._apply_paragraph_format(para)
            self._add_inline_content(para, element)
            return

        # Ordered lists
        if tag_name == "ol":
            for li in element.find_all("li", recursive=False):
                para = document.add_paragraph(style="List Number")
                self._apply_paragraph_format(para)
                self._add_inline_content(para, li)
            return

        # Unordered lists
        if tag_name == "ul":
            for li in element.find_all("li", recursive=False):
                para = document.add_paragraph(style="List Bullet")
                self._apply_paragraph_format(para)
                self._add_inline_content(para, li)
            return

        # Tables
        if tag_name == "table":
            self._process_table(element, document)
            return

        # Images (block-level)
        if tag_name == "img":
            self._process_image(element, document)
            return

        # Div or other containers - recurse into children
        if tag_name in ("div", "section", "article", "span"):
            for child in element.children:
                if isinstance(child, NavigableString):
                    text = child.strip()
                    if text:
                        para = document.add_paragraph()
                        self._apply_paragraph_format(para)
                        para.add_run(text)
                elif isinstance(child, Tag):
                    self._process_element(child, document)
            return

        # Fallback: treat as paragraph with text content
        text = element.get_text(strip=True)
        if text:
            para = document.add_paragraph()
            self._apply_paragraph_format(para)
            self._add_inline_content(para, element)

    def _process_table(self, table_element: Tag, document: Document) -> None:
        """Convert an HTML table to a Word table.

        Args:
            table_element: BeautifulSoup <table> Tag.
            document: python-docx Document to add the table to.
        """
        rows = []

        # Extract header rows
        thead = table_element.find("thead")
        if thead:
            for tr in thead.find_all("tr"):
                cells = []
                for cell in tr.find_all(["th", "td"]):
                    cells.append(cell.get_text(strip=True))
                if cells:
                    rows.append({"cells": cells, "is_header": True})

        # Extract body rows
        tbody = table_element.find("tbody")
        body_source = tbody if tbody else table_element
        for tr in body_source.find_all("tr", recursive=False if tbody else True):
            # Skip rows already captured in thead
            if thead and tr.parent == thead:
                continue
            cells = []
            for cell in tr.find_all(["td", "th"]):
                cells.append(cell.get_text(strip=True))
            if cells:
                rows.append({"cells": cells, "is_header": False})

        if not rows:
            return

        # Determine max columns
        max_cols = max(len(row["cells"]) for row in rows)

        # Create Word table
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply table style
        table.style = "Table Grid"

        # Fill cells
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data["cells"]):
                if col_idx < max_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text

                    # Bold header cells
                    if row_data["is_header"]:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                            paragraph.paragraph_format.space_after = Pt(3)

        # Add a blank paragraph after the table for spacing
        document.add_paragraph()

    def _resolve_image_bytes(self, src: str) -> Optional[io.BytesIO]:
        """Resolve an image source to a BytesIO buffer.

        Handles both base64 data URIs and local file paths under /uploads/.

        Args:
            src: The src attribute value from an <img> tag.

        Returns:
            A BytesIO buffer containing the image data, or None if unresolvable.
        """
        if not src:
            return None

        # Handle base64 data URIs
        if src.startswith("data:"):
            try:
                # Format: data:image/png;base64,<data>
                header, data = src.split(",", 1)
                image_bytes = base64.b64decode(data)
                return io.BytesIO(image_bytes)
            except (ValueError, Exception):
                return None

        # Handle local /uploads/ paths
        if src.startswith("/uploads/"):
            filename = src.replace("/uploads/", "", 1)
            uploads_dir = Path(__file__).parent.parent.parent / "data" / "uploads"
            file_path = uploads_dir / filename
            if file_path.is_file():
                return io.BytesIO(file_path.read_bytes())
            return None

        # Handle relative uploads/ paths (without leading slash)
        if src.startswith("uploads/"):
            filename = src.replace("uploads/", "", 1)
            uploads_dir = Path(__file__).parent.parent.parent / "data" / "uploads"
            file_path = uploads_dir / filename
            if file_path.is_file():
                return io.BytesIO(file_path.read_bytes())
            return None

        return None

    def _process_image(self, element: Tag, document: Document) -> None:
        """Process a block-level <img> element and add it to the document.

        Resolves the image source (base64 or local path) and adds it
        as a picture in the DOCX. Falls back to placeholder text if
        the image cannot be resolved.

        Args:
            element: BeautifulSoup <img> Tag.
            document: python-docx Document to add the image to.
        """
        src = element.get("src", "")
        image_buffer = self._resolve_image_bytes(src)

        if image_buffer:
            try:
                document.add_picture(image_buffer, width=Inches(2))
            except Exception:
                para = document.add_paragraph()
                self._apply_paragraph_format(para)
                alt = element.get("alt", "[Image]")
                para.add_run(f"[{alt}]")
        else:
            para = document.add_paragraph()
            self._apply_paragraph_format(para)
            alt = element.get("alt", "[Image]")
            para.add_run(f"[{alt}]")

    def _add_inline_image(self, paragraph, element: Tag) -> None:
        """Add an inline <img> element to a paragraph.

        Resolves the image source and adds it as an inline shape within
        the paragraph. Falls back to placeholder text if unresolvable.

        Args:
            paragraph: python-docx paragraph to add the image to.
            element: BeautifulSoup <img> Tag.
        """
        src = element.get("src", "")
        image_buffer = self._resolve_image_bytes(src)

        if image_buffer:
            try:
                run = paragraph.add_run()
                run.add_picture(image_buffer, width=Inches(2))
            except Exception:
                alt = element.get("alt", "[Image]")
                run = paragraph.add_run(f"[{alt}]")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        else:
            alt = element.get("alt", "[Image]")
            run = paragraph.add_run(f"[{alt}]")
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def _add_inline_content(self, paragraph, element: Tag) -> None:
        """Add inline content from an element to a paragraph.

        Handles nested <strong>, <b>, <em>, <i>, <u> tags for formatting.

        Args:
            paragraph: python-docx paragraph to add content to.
            element: BeautifulSoup Tag containing inline content.
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
            elif isinstance(child, Tag):
                child_tag = child.name.lower() if child.name else ""

                if child_tag in ("strong", "b"):
                    self._add_inline_with_format(
                        paragraph, child, bold=True
                    )
                elif child_tag in ("em", "i"):
                    self._add_inline_with_format(
                        paragraph, child, italic=True
                    )
                elif child_tag == "u":
                    self._add_inline_with_format(
                        paragraph, child, underline=True
                    )
                elif child_tag == "br":
                    paragraph.add_run("\n")
                elif child_tag == "img":
                    self._add_inline_image(paragraph, child)
                else:
                    # For other inline tags, just get their text
                    self._add_inline_content(paragraph, child)

    def _add_inline_with_format(
        self,
        paragraph,
        element: Tag,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
    ) -> None:
        """Add inline element text with specific formatting.

        Args:
            paragraph: python-docx paragraph to add text to.
            element: BeautifulSoup Tag with text content.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            underline: Whether text should be underlined.
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
            elif isinstance(child, Tag):
                child_tag = child.name.lower() if child.name else ""
                nested_bold = bold or child_tag in ("strong", "b")
                nested_italic = italic or child_tag in ("em", "i")
                nested_underline = underline or child_tag == "u"
                self._add_inline_with_format(
                    paragraph,
                    child,
                    bold=nested_bold,
                    italic=nested_italic,
                    underline=nested_underline,
                )

    def _apply_paragraph_format(self, paragraph) -> None:
        """Apply standard paragraph formatting.

        Sets 1.5 line spacing and 6pt after spacing.

        Args:
            paragraph: python-docx paragraph to format.
        """
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)

    def _apply_heading_font(self, paragraph) -> None:
        """Apply Times New Roman bold font to heading paragraph runs.

        Args:
            paragraph: python-docx heading paragraph.
        """
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.bold = True

    def _strip_tags(self, html: str) -> str:
        """Remove all HTML tags from text.

        Args:
            html: HTML string to strip.

        Returns:
            Plain text without HTML tags.
        """
        clean = re.sub(r"<[^>]+>", "", html)
        # Decode HTML entities
        clean = clean.replace("&amp;", "&")
        clean = clean.replace("&lt;", "<")
        clean = clean.replace("&gt;", ">")
        clean = clean.replace("&quot;", '"')
        clean = clean.replace("&#39;", "'")
        return clean.strip()
