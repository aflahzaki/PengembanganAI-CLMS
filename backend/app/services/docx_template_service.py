"""DOCX Template parsing service for the Template Library.

Scans a directory of .docx template files, parses them into HTML
with smart variable highlighting, and provides CRUD operations
for template management.
"""

import os
import re
import unicodedata
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.config import settings


class DocxTemplateService:
    """Service for managing and parsing DOCX contract templates.

    Provides functionality to:
    - List available templates from the docx directory
    - Parse templates to HTML with smart variable highlighting
    - Extract and classify variables
    - Upload and delete templates
    """

    def __init__(self, templates_dir: Optional[Path] = None):
        """Initialize the service with the templates directory.

        Args:
            templates_dir: Path to the docx templates directory.
                Defaults to settings.templates_absolute_path / "docx".
        """
        if templates_dir is None:
            self.templates_dir = settings.templates_absolute_path / "docx"
        else:
            self.templates_dir = templates_dir
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def _slugify(self, filename: str) -> str:
        """Generate a URL-safe slug from a template filename.

        Strips the date prefix (YYYYMMDD) and 'PLN - BPR -' prefix,
        then converts the remaining name to a lowercase hyphenated slug.

        Args:
            filename: Original docx filename (with extension).

        Returns:
            URL-safe slug string.

        Example:
            "20241007 PLN - BPR - Template Kontrak Lumsum.docx"
            -> "template-kontrak-lumsum"
        """
        # Remove .docx extension
        name = filename.replace(".docx", "").replace(".DOCX", "")

        # Remove date prefix (8 digits followed by space)
        name = re.sub(r"^\d{8}\s*", "", name)

        # Remove "PLN - BPR - " or "PLN - BPR -" prefix
        name = re.sub(r"^PLN\s*-\s*BPR\s*-\s*", "", name)

        # Normalize unicode
        name = unicodedata.normalize("NFKD", name)

        # Convert to lowercase and replace non-alphanumeric with hyphens
        name = name.lower().strip()
        name = re.sub(r"[^a-z0-9]+", "-", name)
        name = name.strip("-")

        return name

    def list_templates(self) -> List[dict]:
        """List all available DOCX templates.

        Returns:
            List of template info dictionaries with id, name, filename,
            variable_count, and file_size_bytes.
        """
        templates = []
        if not self.templates_dir.exists():
            return templates

        for filepath in sorted(self.templates_dir.glob("*.docx")):
            template_id = self._slugify(filepath.name)
            variables = self._extract_variables_from_file(filepath)
            file_size = filepath.stat().st_size

            # Generate display name from slug
            display_name = self._display_name_from_filename(filepath.name)

            templates.append({
                "id": template_id,
                "name": display_name,
                "filename": filepath.name,
                "variable_count": len(variables),
                "variables": variables,
                "file_size_bytes": file_size,
            })

        return templates

    def _display_name_from_filename(self, filename: str) -> str:
        """Generate a human-readable display name from a filename.

        Args:
            filename: Original docx filename.

        Returns:
            Human-readable display name.
        """
        name = filename.replace(".docx", "").replace(".DOCX", "")
        # Remove date prefix
        name = re.sub(r"^\d{8}\s*", "", name)
        # Remove "PLN - BPR - " prefix
        name = re.sub(r"^PLN\s*-\s*BPR\s*-\s*", "", name)
        return name.strip()

    def get_template_by_id(self, template_id: str) -> Optional[Path]:
        """Find a template file by its slug ID.

        Args:
            template_id: The URL-safe slug identifier.

        Returns:
            Path to the template file, or None if not found.
        """
        if not self.templates_dir.exists():
            return None

        for filepath in self.templates_dir.glob("*.docx"):
            if self._slugify(filepath.name) == template_id:
                return filepath
        return None

    def parse_to_html(self, template_id: str) -> Optional[dict]:
        """Parse a DOCX template to HTML with smart variable highlighting.

        Args:
            template_id: The URL-safe slug identifier.

        Returns:
            Dictionary with id, name, html_content, and variables,
            or None if template not found.
        """
        filepath = self.get_template_by_id(template_id)
        if filepath is None:
            return None

        doc = Document(str(filepath))
        html_parts = []

        for element in self._iter_document_elements(doc):
            if isinstance(element, Paragraph):
                html_parts.append(self._paragraph_to_html(element))
            elif isinstance(element, Table):
                html_parts.append(self._table_to_html(element))

        raw_html = "\n".join(html_parts)
        highlighted_html = self._apply_variable_highlighting(raw_html)

        variables = self._extract_variables_from_file(filepath)

        return {
            "id": template_id,
            "name": self._display_name_from_filename(filepath.name),
            "html_content": highlighted_html,
            "variables": variables,
        }

    def _iter_document_elements(self, doc: DocxDocument):
        """Iterate through document elements (paragraphs and tables) in order.

        Args:
            doc: python-docx Document object.

        Yields:
            Paragraph or Table objects in document order.
        """
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                yield Paragraph(element, doc)
            elif tag == "tbl":
                yield Table(element, doc)

    def _paragraph_to_html(self, paragraph: Paragraph) -> str:
        """Convert a paragraph to an HTML string.

        Maps paragraph styles to appropriate HTML tags:
        - Heading 1-9 -> <h1>-<h6>
        - List items -> <li> wrapped in appropriate list tags
        - Normal/Body -> <p>

        Args:
            paragraph: python-docx Paragraph object.

        Returns:
            HTML string representation.
        """
        text = self._runs_to_html(paragraph)
        if not text.strip():
            return ""

        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""

        # Headings
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
                level = min(level, 6)
            except ValueError:
                level = 2
            return f"<h{level}>{text}</h{level}>"

        # List items
        if "list" in style_name or paragraph._element.pPr is not None and paragraph._element.pPr.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        ) is not None:
            return f"<li>{text}</li>"

        # Default: paragraph
        return f"<p>{text}</p>"

    def _runs_to_html(self, paragraph: Paragraph) -> str:
        """Convert paragraph runs to HTML with inline formatting.

        Uses a two-pass approach: first concatenates all run text to find
        bracket boundaries, then splits formatting around complete bracket
        expressions to ensure highlighting regex can match them.

        Args:
            paragraph: python-docx Paragraph object.

        Returns:
            HTML string with inline formatting tags.
        """
        if not paragraph.runs:
            return ""

        # Collect run info: (text, bold, italic, underline)
        run_info = []
        for run in paragraph.runs:
            if run.text:
                run_info.append((run.text, run.bold, run.italic, run.underline))

        if not run_info:
            return ""

        # Concatenate all text to find bracket boundaries
        full_text = "".join(r[0] for r in run_info)

        # Find all bracket spans in the concatenated text
        bracket_spans = []
        pattern = re.compile(r"\[([^\[\]]+)\]")
        for m in pattern.finditer(full_text):
            bracket_spans.append((m.start(), m.end()))

        # Build character-level formatting map
        char_formats = []
        offset = 0
        for text, bold, italic, underline in run_info:
            for _ in text:
                char_formats.append((bold, italic, underline))
            offset += len(text)

        # Build output segments: for text inside brackets, emit the bracket
        # content without formatting wrapping (formatting will be inside the
        # bracket text). For text outside brackets, apply per-run formatting.
        parts = []
        pos = 0
        bracket_idx = 0

        while pos < len(full_text):
            # Check if we're at the start of a bracket span
            if bracket_idx < len(bracket_spans) and pos == bracket_spans[bracket_idx][0]:
                start, end = bracket_spans[bracket_idx]
                bracket_text = full_text[start:end]
                # Escape HTML in the bracket text
                bracket_text = bracket_text.replace("&", "&amp;")
                bracket_text = bracket_text.replace("<", "&lt;")
                bracket_text = bracket_text.replace(">", "&gt;")
                parts.append(bracket_text)
                pos = end
                bracket_idx += 1
            else:
                # Find the end of this non-bracket segment
                next_bracket_start = bracket_spans[bracket_idx][0] if bracket_idx < len(bracket_spans) else len(full_text)
                segment_end = next_bracket_start

                # Process character by character with formatting
                # Group consecutive characters with same formatting
                i = pos
                while i < segment_end:
                    fmt = char_formats[i]
                    # Find extent of same formatting
                    j = i
                    while j < segment_end and char_formats[j] == fmt:
                        j += 1
                    segment_text = full_text[i:j]
                    # Escape HTML
                    segment_text = segment_text.replace("&", "&amp;")
                    segment_text = segment_text.replace("<", "&lt;")
                    segment_text = segment_text.replace(">", "&gt;")
                    # Apply formatting
                    bold, italic, underline = fmt
                    if bold:
                        segment_text = f"<strong>{segment_text}</strong>"
                    if italic:
                        segment_text = f"<em>{segment_text}</em>"
                    if underline:
                        segment_text = f"<u>{segment_text}</u>"
                    parts.append(segment_text)
                    i = j

                pos = segment_end

        return "".join(parts)

    def _table_to_html(self, table: Table) -> str:
        """Convert a table to an HTML table string.

        Args:
            table: python-docx Table object.

        Returns:
            HTML table string.
        """
        rows_html = []
        for row in table.rows:
            cells_html = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Escape HTML entities
                cell_text = cell_text.replace("&", "&amp;")
                cell_text = cell_text.replace("<", "&lt;")
                cell_text = cell_text.replace(">", "&gt;")
                cells_html.append(f"<td>{cell_text}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")

        return f"<table>{''.join(rows_html)}</table>"

    def _apply_variable_highlighting(self, html: str) -> str:
        """Apply smart variable highlighting to HTML content.

        Three types of highlighting based on bracket content:
        1. Ellipsis content ([...], [...], [(...)])
           -> <mark> with bright yellow #FFEB3B (dynamic/mandatory)
        2. Short descriptive (<=60 chars)
           -> <mark> with light yellow #FFF9C4 (editable)
        3. Long text (>60 chars)
           -> <span> with gray italic (instruction, no highlight)

        Args:
            html: HTML content with bracketed variables.

        Returns:
            HTML with variable highlighting applied.
        """
        pattern = r"\[([^\[\]]+)\]"

        def highlight_match(match: re.Match) -> str:
            content = match.group(1)
            full_match = match.group(0)
            var_type = self._classify_variable(content)

            if var_type == "dynamic":
                return (
                    f'<mark style="background-color: #FFEB3B; '
                    f'padding: 2px 4px; border-radius: 2px;">'
                    f'{full_match}</mark>'
                )
            elif var_type == "editable":
                return (
                    f'<mark style="background-color: #FFF9C4; '
                    f'padding: 2px 4px; border-radius: 2px;">'
                    f'{full_match}</mark>'
                )
            elif var_type == "instruction":
                return (
                    f'<span style="color: gray; font-style: italic;">'
                    f'{full_match}</span>'
                )
            else:
                return full_match

        return re.sub(pattern, highlight_match, html)

    def _classify_variable(self, content: str) -> str:
        """Classify a variable based on its bracketed content.

        Classification rules:
        - "dynamic": content is ONLY dots/ellipsis characters
          (e.g., "...", "...", "(...)").
        - "editable": content is <= 60 characters (short descriptive text).
        - "instruction": content is > 60 characters (long instruction text).

        Args:
            content: The text inside brackets (without brackets).

        Returns:
            One of "dynamic", "editable", or "instruction".
        """
        # Check if content is only dots, ellipsis chars, parentheses, and whitespace
        # This covers: "...", "...", "(..)", "(...)", "(  ...  )", etc.
        stripped = content.strip()
        # Remove parentheses and whitespace for the check
        check_content = stripped.replace("(", "").replace(")", "").replace(" ", "")

        # Check if all remaining characters are dot-like
        if check_content and all(c in ".\u2026\u2025\u00b7" for c in check_content):
            return "dynamic"

        # Check length for editable vs instruction
        if len(content) > 60:
            return "instruction"

        return "editable"

    def _extract_variables_from_file(self, filepath: Path) -> List[dict]:
        """Extract and classify all variables from a DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            List of variable dictionaries with name, type, and full_text.
        """
        doc = Document(str(filepath))
        variables = []
        seen = set()

        # Extract from all paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            self._extract_vars_from_text(text, variables, seen)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._extract_vars_from_text(cell.text, variables, seen)

        return variables

    def _extract_vars_from_text(
        self, text: str, variables: List[dict], seen: set
    ) -> None:
        """Extract variables from a text string and add to the list.

        Args:
            text: Text to search for variables.
            variables: List to append found variables to.
            seen: Set of already-seen variable texts (for deduplication).
        """
        pattern = r"\[([^\[\]]+)\]"
        for match in re.finditer(pattern, text):
            content = match.group(1)
            if content in seen:
                continue
            seen.add(content)

            var_type = self._classify_variable(content)
            variables.append({
                "name": content if var_type != "dynamic" else "...",
                "type": var_type,
                "full_text": match.group(0),
            })

    def save_uploaded_template(self, filename: str, content: bytes) -> dict:
        """Save an uploaded DOCX file to the templates directory.

        Args:
            filename: Original filename of the uploaded file.
            content: Raw bytes of the DOCX file.

        Returns:
            Template info dictionary for the saved file.

        Raises:
            ValueError: If the file is not a .docx file or is not a valid DOCX.
        """
        if not filename.lower().endswith(".docx"):
            raise ValueError("Only .docx files are accepted")

        # Sanitize filename to prevent path traversal
        safe_filename = Path(filename).name
        if not safe_filename or safe_filename.startswith("."):
            raise ValueError("Invalid filename")

        # Validate that the content is actually a valid DOCX file
        from io import BytesIO
        try:
            Document(BytesIO(content))
        except Exception:
            raise ValueError("File is not a valid DOCX document")

        filepath = self.templates_dir / safe_filename
        filepath.write_bytes(content)

        template_id = self._slugify(safe_filename)
        variables = self._extract_variables_from_file(filepath)
        file_size = filepath.stat().st_size
        display_name = self._display_name_from_filename(safe_filename)

        return {
            "id": template_id,
            "name": display_name,
            "filename": safe_filename,
            "variable_count": len(variables),
            "variables": variables,
            "file_size_bytes": file_size,
        }

    def delete_template(self, template_id: str) -> bool:
        """Delete a template file by its slug ID.

        Args:
            template_id: The URL-safe slug identifier.

        Returns:
            True if deleted, False if not found.
        """
        filepath = self.get_template_by_id(template_id)
        if filepath is None:
            return False

        filepath.unlink()
        return True
