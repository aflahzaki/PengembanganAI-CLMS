"""
Script to generate Panduan_AI_CLMS.docx using python-docx.
Run: python docs/generate_panduan.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set borders for the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_heading1(doc, text):
    """Add a Heading 1 style paragraph."""
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 1']
    run = heading.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    set_paragraph_spacing(heading, before=18, after=6)
    return heading


def add_heading2(doc, text):
    """Add a Heading 2 style paragraph."""
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 2']
    run = heading.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 70, 127)
    set_paragraph_spacing(heading, before=12, after=6)
    return heading


def add_body(doc, text, bold=False):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = bold
    set_paragraph_spacing(para, before=3, after=3, line_spacing=1.5)
    return para


def add_bullet(doc, text, level=0):
    """Add a bullet list item."""
    para = doc.add_paragraph(style='List Bullet')
    para.clear()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    set_paragraph_spacing(para, before=2, after=2, line_spacing=1.5)
    return para


def add_numbered(doc, text, level=0):
    """Add a numbered list item."""
    para = doc.add_paragraph(style='List Number')
    para.clear()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    set_paragraph_spacing(para, before=2, after=2, line_spacing=1.5)
    return para


def create_table(doc, headers, rows):
    """Create a formatted table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "003366")
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F2F7FC")

    doc.add_paragraph()  # spacing after table
    return table


def generate_document():
    """Generate the full Panduan AI CLMS document."""
    doc = Document()

    # Set default margins (2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=72, after=12)
    run = title.add_run("PANDUAN AI CLMS")
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=6, after=6)
    run = subtitle.add_run("Contract Lifecycle Management System")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 70, 127)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(tagline, before=24, after=72)
    run = tagline.add_run("Untuk Pemula yang Baru Mengenal AI")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.italic = True

    # =========================================================================
    # SECTION 1: Apa itu AI CLMS?
    # =========================================================================
    add_heading1(doc, "1. Apa itu AI CLMS?")

    add_body(doc, (
        "AI CLMS (Contract Lifecycle Management System) adalah sistem yang menggunakan "
        "kecerdasan buatan (Artificial Intelligence/AI) untuk membantu membuat dokumen kontrak "
        "PLN secara otomatis."
    ))

    add_body(doc, (
        "Penting untuk dipahami: sistem ini BUKAN menggantikan manusia. AI CLMS adalah alat "
        "bantu yang mempercepat pekerjaan pembuatan draft kontrak yang biasanya memakan waktu "
        "lama jika dikerjakan secara manual. Hasil dari AI tetap perlu direview dan disetujui "
        "oleh manusia sebelum digunakan secara resmi."
    ))

    add_body(doc, (
        "Dengan AI CLMS, proses yang biasanya membutuhkan waktu berjam-jam untuk menyusun "
        "draft kontrak bisa dipercepat menjadi hitungan menit. User cukup memasukkan informasi "
        "dasar seperti nama pekerjaan, nilai kontrak, dan jangka waktu, lalu sistem akan "
        "menghasilkan draft kontrak yang siap untuk direview."
    ))

    # =========================================================================
    # SECTION 2: Bagaimana AI Ini Bekerja?
    # =========================================================================
    add_heading1(doc, "2. Bagaimana AI Ini Bekerja? (Penjelasan Sederhana)")

    add_heading2(doc, "Analogi Sederhana")

    add_body(doc, "Bayangkan AI seperti seorang asisten pintar yang sudah membaca ratusan contoh kontrak:")

    add_bullet(doc, (
        'AI seperti "asisten pintar" yang sudah membaca banyak contoh kontrak sebelumnya. '
        "Dia sudah paham pola, struktur, dan bahasa hukum yang biasa digunakan."
    ))
    add_bullet(doc, (
        "Saat diminta membuat kontrak baru, AI menyusun kalimat berdasarkan pola yang sudah "
        "dipelajari dari contoh-contoh sebelumnya."
    ))
    add_bullet(doc, (
        'AI menggunakan "template" (cetakan) yang sudah ada, lalu mengisi variabel sesuai '
        "kebutuhan user. Mirip seperti mengisi formulir, tapi lebih cerdas karena bisa "
        "menyesuaikan kalimat."
    ))

    add_heading2(doc, "2 Mode Kerja Sistem")

    add_body(doc, "Mode 1: Dari Template (Tanpa AI)", bold=True)
    add_body(doc, (
        "Pada mode ini, AI TIDAK digunakan sama sekali. Sistem hanya menampilkan template "
        "kontrak yang sudah ada. User tinggal mengisi bagian-bagian yang di-highlight kuning "
        "(variabel/placeholder) dengan informasi yang sesuai. Mode ini cocok jika Anda sudah "
        "tahu persis jenis kontrak yang dibutuhkan."
    ))

    add_body(doc, "Mode 2: Generate Baru (Dengan AI Groq)", bold=True)
    add_body(doc, (
        "Pada mode ini, AI (Groq) digunakan untuk menyusun dokumen kontrak baru dari nol "
        "berdasarkan deskripsi yang diberikan user. User cukup menjelaskan kebutuhan kontrak "
        "dalam bahasa sehari-hari, dan AI akan menghasilkan draft kontrak lengkap. Mode ini "
        "cocok jika Anda butuh draft cepat atau kontrak dengan kebutuhan khusus."
    ))

    # =========================================================================
    # SECTION 3: Komponen Utama Sistem
    # =========================================================================
    add_heading1(doc, "3. Komponen Utama Sistem")

    add_body(doc, "Sistem AI CLMS terdiri dari beberapa komponen utama yang bekerja sama:")

    components = [
        ["Komponen", "Teknologi", "Fungsi"],
    ]
    component_rows = [
        ["Frontend (Website)", "SvelteKit", "Tampilan yang dilihat dan digunakan user. Di sinilah user memilih template, mengisi data, dan melihat hasil."],
        ["Backend (Server API)", "FastAPI (Python)", '"Otak" sistem yang memproses semua permintaan. Menerima input dari frontend, berkomunikasi dengan AI, dan mengirim hasil kembali.'],
        ["LLM/AI (Groq)", "Llama 3.3 70B", "Mesin AI yang menghasilkan teks kontrak. Kita menggunakan Groq yang menyediakan layanan AI secara gratis."],
        ["Database (ChromaDB)", "ChromaDB", "Tempat menyimpan data template kontrak dan klausul dalam bentuk vektor untuk pencarian cerdas."],
        ["TipTap Editor", "TipTap", "Editor dokumen di browser (seperti Google Docs versi mini). User bisa mengedit hasil draft di sini sebelum export."],
    ]

    create_table(doc, components[0], component_rows)

    # =========================================================================
    # SECTION 4: Istilah-Istilah Penting
    # =========================================================================
    add_heading1(doc, "4. Istilah-Istilah Penting")

    add_body(doc, (
        "Berikut adalah istilah-istilah yang sering muncul dalam pengembangan dan penggunaan "
        "AI CLMS. Jangan khawatir jika belum paham semuanya, pahami sedikit demi sedikit:"
    ))

    terms_headers = ["Istilah", "Penjelasan Sederhana"]
    terms_rows = [
        ["API (Application Programming Interface)",
         '"Jembatan" antara frontend dan backend. API memungkinkan dua sistem berbicara satu sama lain.'],
        ["Endpoint",
         "Alamat spesifik di API untuk fungsi tertentu. Contoh: /api/draft untuk membuat draft kontrak."],
        ["LLM (Large Language Model)",
         "Model AI besar yang bisa memahami dan menghasilkan teks seperti manusia. Contoh: GPT, Llama, Gemini."],
        ["Token",
         "Satuan teks terkecil yang dipahami AI. Kira-kira 1 kata sama dengan 1-2 token."],
        ["Prompt",
         "Instruksi atau perintah yang diberikan ke AI. Semakin jelas prompt-nya, semakin bagus hasil AI."],
        ["Template",
         "Cetakan dokumen yang sudah jadi strukturnya, tinggal diisi variabel/data yang sesuai."],
        ["Variable/Placeholder",
         "Bagian di template yang perlu diisi oleh user. Di sistem ini ditandai dengan highlight kuning. Contoh: [Nama Pekerjaan]."],
        ["RAG (Retrieval-Augmented Generation)",
         "Teknik AI yang mengambil data relevan dari database terlebih dahulu sebelum men-generate teks. Hasilnya lebih akurat."],
        ["ChromaDB",
         "Database vektor khusus untuk menyimpan dan mencari klausul kontrak berdasarkan kemiripan makna."],
        ["DOCX",
         "Format file Microsoft Word. Ini adalah format output utama dari sistem AI CLMS."],
        ["HTML",
         "Bahasa markup untuk menampilkan konten di website. Editor TipTap menggunakan HTML."],
        ["Deployment",
         "Proses memasang/menginstall aplikasi ke server agar bisa diakses oleh orang lain melalui internet."],
        ["Docker",
         'Alat untuk "membungkus" aplikasi beserta semua kebutuhannya agar mudah di-deploy di server manapun.'],
        ["Git/GitHub",
         "Sistem penyimpanan kode dan kolaborasi. Git mencatat semua perubahan, GitHub adalah platform online-nya."],
        ["Pull Request",
         "Permintaan untuk menggabungkan perubahan kode yang sudah dibuat ke kode utama (main branch)."],
        ["Environment Variable (.env)",
         "File konfigurasi rahasia yang berisi API key dan pengaturan sensitif lainnya. Tidak boleh di-share."],
    ]

    create_table(doc, terms_headers, terms_rows)

    # =========================================================================
    # SECTION 5: Alur Kerja Sistem
    # =========================================================================
    add_heading1(doc, "5. Alur Kerja Sistem (Step by Step)")

    add_body(doc, "Berikut adalah alur lengkap dari user membuka website sampai mendapatkan dokumen kontrak:")

    add_heading2(doc, "Alur Umum")

    add_numbered(doc, "User membuka website AI CLMS melalui browser (alamat: localhost:5173 untuk development)")
    add_numbered(doc, "User memilih mode: Dari Template atau Generate Baru")

    add_heading2(doc, "Alur Mode Template")

    add_numbered(doc, "User memilih template kontrak yang diinginkan dari daftar yang tersedia")
    add_numbered(doc, "Template ditampilkan di editor TipTap dengan bagian variabel di-highlight kuning")
    add_numbered(doc, "User mengisi semua variabel kuning dengan data yang sesuai (nama pekerjaan, nilai, dll)")
    add_numbered(doc, "User bisa mengedit teks lainnya jika perlu")
    add_numbered(doc, "Klik Export ke DOCX untuk download file Word yang sudah jadi")

    add_heading2(doc, "Alur Mode Generate (AI)")

    add_numbered(doc, "User mengisi deskripsi pekerjaan dan variabel yang diperlukan")
    add_numbered(doc, "Sistem mengirim data ke backend API")
    add_numbered(doc, "Backend memproses request dan mengirim prompt ke Groq AI")
    add_numbered(doc, "AI menghasilkan teks kontrak berdasarkan deskripsi yang diberikan")
    add_numbered(doc, "Hasil muncul di editor TipTap untuk direview dan diedit")
    add_numbered(doc, "User mereview dan memperbaiki jika ada yang perlu diubah")
    add_numbered(doc, "Klik Export ke DOCX untuk download file Word")

    # =========================================================================
    # SECTION 6: Tentang Groq API
    # =========================================================================
    add_heading1(doc, "6. Tentang Groq API")

    add_body(doc, (
        "Groq adalah penyedia layanan AI yang memberikan akses gratis ke model-model AI canggih. "
        "Berikut informasi penting tentang Groq yang digunakan di sistem ini:"
    ))

    add_bullet(doc, "Model yang dipakai: Llama 3.3 70B (model AI dengan 70 miliar parameter)")
    add_bullet(doc, "Biaya: GRATIS (ada batasan penggunaan)")
    add_bullet(doc, "Batasan: 30 request per menit (lebih dari cukup untuk penggunaan normal)")
    add_bullet(doc, "Kecepatan: Sangat cepat, sekitar 394 token per detik")
    add_bullet(doc, "Kualitas: Setara dengan model berbayar untuk tugas pembuatan kontrak")

    add_heading2(doc, "Cara Mendaftar Groq")

    add_numbered(doc, "Buka website: console.groq.com")
    add_numbered(doc, "Klik Sign Up dan daftar menggunakan email atau akun Google")
    add_numbered(doc, "Setelah masuk, buka menu API Keys")
    add_numbered(doc, 'Klik "Create API Key" dan beri nama (contoh: clms-dev)')
    add_numbered(doc, "Copy API key yang muncul (PENTING: simpan baik-baik, hanya muncul sekali!)")
    add_numbered(doc, "Paste API key ke file backend/.env pada variabel GROQ_API_KEY")

    # =========================================================================
    # SECTION 7: Tentang Template Kontrak
    # =========================================================================
    add_heading1(doc, "7. Tentang Template Kontrak")

    add_body(doc, "Sistem AI CLMS menyediakan 8 jenis template kontrak PLN yang siap pakai:")

    add_numbered(doc, "Kontrak Harga Satuan - untuk pekerjaan yang dihitung per satuan volume")
    add_numbered(doc, "Kontrak Lumsum - untuk pekerjaan dengan harga tetap/borongan")
    add_numbered(doc, "Kontrak Gabungan - kombinasi harga satuan dan lumsum")
    add_numbered(doc, "Kontrak Terima Jadi (Turnkey) - untuk proyek yang diserahkan dalam kondisi siap pakai")
    add_numbered(doc, "Kontrak Kesepakatan Harga Satuan - harga satuan yang disepakati bersama")
    add_numbered(doc, "Kontrak Payung - kontrak jangka panjang untuk kebutuhan berulang")
    add_numbered(doc, "Kontrak Waktu Penugasan - untuk jasa berdasarkan waktu kerja")
    add_numbered(doc, "Kontrak Biaya Plus Imbalan - biaya aktual ditambah fee/imbalan")

    add_heading2(doc, "Tentang Variabel Template")

    add_bullet(doc, "Setiap template memiliki variabel (placeholder) yang perlu diisi oleh user")
    add_bullet(doc, "Variabel ditandai dengan highlight kuning di editor agar mudah terlihat")
    add_bullet(doc, "Contoh variabel: [Nama Pekerjaan], [Nilai Kontrak], [Jangka Waktu], dll")
    add_bullet(doc, "User bisa meng-upload template baru sesuai kebutuhan")
    add_bullet(doc, "Template yang tidak diperlukan bisa dihapus dari sistem")

    # =========================================================================
    # SECTION 8: Cara Menjalankan Sistem
    # =========================================================================
    add_heading1(doc, "8. Cara Menjalankan Sistem")

    add_body(doc, (
        "Berikut langkah-langkah untuk menjalankan sistem AI CLMS di komputer lokal "
        "(untuk development/pengembangan):"
    ))

    add_heading2(doc, "Menjalankan Backend")

    add_numbered(doc, "Buka terminal (Command Prompt / Terminal)")
    add_numbered(doc, "Masuk ke folder backend: cd backend")
    add_numbered(doc, "Aktifkan virtual environment: source venv/bin/activate (Linux/Mac) atau venv\\Scripts\\activate (Windows)")
    add_numbered(doc, "Jalankan server: uvicorn app.main:app --reload --port 8000")
    add_numbered(doc, "Backend akan berjalan di: http://localhost:8000")

    add_heading2(doc, "Menjalankan Frontend")

    add_numbered(doc, "Buka terminal baru (jangan tutup terminal backend)")
    add_numbered(doc, "Masuk ke folder frontend: cd frontend")
    add_numbered(doc, "Install dependencies (pertama kali saja): npm install")
    add_numbered(doc, "Jalankan dev server: npm run dev")
    add_numbered(doc, "Frontend akan berjalan di: http://localhost:5173")

    add_heading2(doc, "Mengakses Aplikasi")

    add_body(doc, (
        "Buka browser (Chrome/Firefox/Edge) dan ketik: http://localhost:5173. "
        "Pastikan backend sudah jalan terlebih dahulu sebelum menggunakan fitur AI Generate."
    ))

    # =========================================================================
    # SECTION 9: File-File Penting
    # =========================================================================
    add_heading1(doc, "9. File-File Penting")

    add_body(doc, "Berikut adalah file-file penting dalam project AI CLMS beserta fungsinya:")

    files_headers = ["File/Folder", "Fungsi"]
    files_rows = [
        ["backend/.env", "File konfigurasi rahasia. Berisi API key Groq, URL database, dan pengaturan lainnya."],
        ["backend/app/main.py", "Entry point (titik masuk) backend. File pertama yang dijalankan saat server start."],
        ["backend/app/services/llm_service.py", "Kode yang berkomunikasi langsung dengan AI Groq. Di sinilah prompt dikirim dan respons diterima."],
        ["backend/data/templates/docx/", "Folder yang berisi semua file template kontrak dalam format DOCX."],
        ["frontend/src/routes/", "Folder yang berisi halaman-halaman website (routing SvelteKit)."],
        ["postman/", "File koleksi Postman untuk testing API secara manual. Berguna untuk debugging."],
    ]

    create_table(doc, files_headers, files_rows)

    # =========================================================================
    # SECTION 10: FAQ
    # =========================================================================
    add_heading1(doc, "10. FAQ (Pertanyaan yang Sering Diajukan)")

    # Q1
    add_body(doc, "Q: Apakah AI ini bisa salah?", bold=True)
    add_body(doc, (
        "A: Ya, AI tidak sempurna dan bisa menghasilkan teks yang kurang tepat. Oleh karena itu, "
        "hasil AI SELALU perlu direview oleh manusia sebelum digunakan. Jangan langsung pakai "
        "hasil AI tanpa dicek terlebih dahulu."
    ))

    # Q2
    add_body(doc, "Q: Apakah data kontrak aman?", bold=True)
    add_body(doc, (
        "A: Data disimpan di server sendiri (lokal). AI Groq hanya menerima teks prompt secara "
        "sementara untuk diproses dan tidak menyimpan data Anda secara permanen. Namun, untuk "
        "data sangat sensitif, konsultasikan dengan tim IT terlebih dahulu."
    ))

    # Q3
    add_body(doc, "Q: Bagaimana jika Groq sedang down/tidak bisa diakses?", bold=True)
    add_body(doc, (
        "A: Sistem memiliki mekanisme fallback otomatis. Jika Groq tidak tersedia, sistem akan "
        "beralih ke mode template-fill (tanpa AI). Anda tetap bisa menggunakan template yang "
        "ada dan mengisi variabel secara manual."
    ))

    # Q4
    add_body(doc, "Q: Apakah perlu koneksi internet?", bold=True)
    add_body(doc, (
        "A: Ya, koneksi internet diperlukan untuk mengakses Groq API (mode Generate). Namun, "
        "mode Template bisa digunakan secara offline karena template sudah tersimpan di server lokal."
    ))

    # Q5
    add_body(doc, "Q: Berapa biaya untuk menggunakan sistem ini?", bold=True)
    add_body(doc, (
        "A: Groq API gratis digunakan. Biaya yang diperlukan hanya untuk server hosting jika "
        "ingin di-deploy agar bisa diakses oleh banyak orang. Untuk development di komputer "
        "lokal, tidak ada biaya sama sekali."
    ))

    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("--- Dokumen ini dibuat sebagai panduan untuk tim pengembang AI CLMS ---")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Panduan_AI_CLMS.docx")
    doc.save(output_path)
    print(f"Document generated successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_document()
